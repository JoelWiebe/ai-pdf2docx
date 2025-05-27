# File: phase2_json_to_docx.py
import json
import argparse
import os
from docx import Document


def parse_markdown_table(markdown_table_content: str):
    lines = [line.strip() for line in markdown_table_content.strip().split('\n') if line.strip()]
    if not lines: return []
    def clean_cell(cell_text): return cell_text.strip()
    table_data = []; header_parsed = False; num_cols = 0
    for i, line in enumerate(lines):
        if not line.startswith('|') or not line.endswith('|'):
            if i == 1 and all(c in '-|: ' for c in line): continue
            continue
        cells = [clean_cell(cell) for cell in line.strip('|').split('|')]
        if not header_parsed:
            table_data.append(cells); num_cols = len(cells); header_parsed = True
        elif len(cells) == num_cols: table_data.append(cells)
        elif len(cells) > num_cols and num_cols > 0:
            table_data.append(cells[:num_cols])
        elif len(cells) < num_cols and num_cols > 0:
            cells.extend([''] * (num_cols - len(cells))); table_data.append(cells)
        elif num_cols == 0 and cells:
            table_data.append(cells); num_cols = len(cells)
    return table_data


def create_docx_from_structured_data(json_content_str: str, output_docx_path: str):
    """Creates a DOCX file from the structured JSON data string using default Word styles."""
    try:
        data = json.loads(json_content_str)
        if "document_elements" not in data or not isinstance(data["document_elements"], list):
            print(f"Error: 'document_elements' key not found or not a list in JSON for {output_docx_path}.")
            return False
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON from file for {output_docx_path}: {e}")
        return False

    doc = Document()

    for element in data.get("document_elements", []):
        el_type = element.get("type")
        content = element.get("content", "")

        if el_type == "heading_1":
            doc.add_paragraph(content, style='Heading 1') 
        elif el_type == "heading_2":
            doc.add_paragraph(content, style='Heading 2') 
        elif el_type == "heading_3":
            doc.add_paragraph(content, style='Heading 3') 
        elif el_type == "paragraph":
            processed_content = ' '.join(content.splitlines()).strip()
            if processed_content:
                doc.add_paragraph(processed_content, style='Normal') 
        elif el_type == "table_markdown":
            if content and content.strip():
                table_data = parse_markdown_table(content)
                if table_data and table_data[0] and len(table_data[0]) > 0:
                    num_cols = len(table_data[0])
                    # Add a paragraph with default spacing before table, if desired,
                    # or manage spacing via table properties if python-docx supports it easily.
                    # For simplicity, Word's default paragraph spacing will apply here.
                    # doc.add_paragraph() # Optional: for explicit spacing before table

                    try:
                        table_docx = doc.add_table(rows=0, cols=num_cols)
                        table_docx.style = 'Table Grid' # 'Table Grid' is a common built-in table style

                        hdr_cells = table_docx.add_row().cells
                        for i, header_text in enumerate(table_data[0]):
                            if i < num_cols:
                                hdr_cells[i].text = header_text
                                if hdr_cells[i].paragraphs and hdr_cells[i].paragraphs[0].runs:
                                    hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                        
                        for row_data in table_data[1:]:
                            if len(row_data) == num_cols:
                                row_cells = table_docx.add_row().cells
                                for i, cell_text in enumerate(row_data):
                                    row_cells[i].text = cell_text
                        
                        # Optional: for explicit spacing after table
                        # doc.add_paragraph()
                    except Exception as e:
                        print(f"Error creating table in DOCX for {output_docx_path}: {e}")
                        # Fallback: add Markdown as text using 'Normal' style
                        doc.add_paragraph(f"[Fallback: Error rendering table. Raw Markdown:]\n{content}", style='Normal')
            # else: empty table content
        else:
            if content:
                 doc.add_paragraph(f"[Unknown type: {el_type}] {content}", style='Normal') # MODIFIED

    try:
        doc.save(output_docx_path)
        print(f"Successfully created DOCX: {output_docx_path}")
        return True
    except Exception as e:
        print(f"Error saving DOCX file {output_docx_path}: {e}")
        return False

def run_json_to_docx_conversion(args):
    """
    Main logic for Phase 2: JSON to DOCX conversion.
    Takes an argparse Namespace object.
    """
    if not os.path.isdir(args.input_json_directory):
        print(f"Error: Input JSON directory '{args.input_json_directory}' not found.")
        return
    
    os.makedirs(args.output_docx_directory, exist_ok=True)

    for filename in os.listdir(args.input_json_directory):
        if filename.lower().endswith(".json"):
            json_path = os.path.join(args.input_json_directory, filename)
            docx_filename = os.path.splitext(filename)[0] + ".docx"
            output_docx_path = os.path.join(args.output_docx_directory, docx_filename)

            if os.path.exists(output_docx_path) and not args.overwrite:
                print(f"Skipping {json_path}, DOCX output already exists: {output_docx_path}")
                continue
            
            print(f"\nProcessing {json_path}...")
            try:
                with open(json_path, "r", encoding="utf-8") as f:
                    json_content_str = f.read()
            except Exception as e:
                print(f"Error reading JSON file {json_path}: {e}")
                continue
            
            if json_content_str.strip():
                create_docx_from_structured_data(json_content_str, output_docx_path)
            else:
                print(f"Skipping empty JSON file: {json_path}")
                
    print("\nPhase 2 (JSON to DOCX) processing complete.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Standalone Phase 2: JSON to DOCX.")
    parser.add_argument("input_json_directory", help="Path to the directory containing JSON files.")
    parser.add_argument("output_docx_directory", help="Path to the directory where DOCX files will be saved.")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite existing DOCX files.")
    args = parser.parse_args()
    run_json_to_docx_conversion(args)